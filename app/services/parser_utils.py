import re
from typing import Dict, List
from io import BytesIO
import PyPDF2

# Robust spacy import with fallback
try:
    import spacy
    try:
        nlp = spacy.load("en_core_web_sm")
        SPACY_AVAILABLE = True
    except OSError:
        print("Warning: spacy model 'en_core_web_sm' not found. Install with: python -m spacy download en_core_web_sm")
        nlp = None
        SPACY_AVAILABLE = False
except ImportError:
    print("Warning: spacy not available. Install with: pip install spacy")
    nlp = None
    SPACY_AVAILABLE = False

# Robust docx import with conflict resolution
try:
    # First try to import python-docx (the correct package)
    from docx import Document
    import docx
    DOCX_AVAILABLE = True
    print("Successfully imported python-docx")
except ImportError:
    try:
        # If that fails, try the old docx package
        import docx
        from docx import Document
        DOCX_AVAILABLE = True
        print("Warning: Using old docx package. Consider upgrading to python-docx")
    except ImportError:
        print("Warning: docx package not found. Install with: pip uninstall docx && pip install python-docx")
        DOCX_AVAILABLE = False
        
        # Create a dummy Document class to prevent errors
        class Document:
            def __init__(self, *args, **kwargs):
                self.paragraphs = []
            
        docx = type('docx', (), {'Document': Document})()

# Regex patterns for text extraction
EMAIL_REGEX = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
PHONE_REGEX = r"(\+?\d{1,3})?[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
LINKEDIN_REGEX = r"https?://(www\.)?linkedin\.com/in/[A-Za-z0-9_-]+/?"
GITHUB_REGEX = r"https?://(www\.)?github\.com/[A-Za-z0-9_-]+/?"

def extract_text(file: BytesIO, filename: str) -> str:
    """Extract text from PDF or DOCX files with error handling"""
    try:
        if filename.endswith('.pdf'):
            reader = PyPDF2.PdfReader(file)
            return "\n".join([page.extract_text() or '' for page in reader.pages])
        elif filename.endswith('.docx'):
            if not DOCX_AVAILABLE:
                raise ValueError("python-docx not available. Cannot process DOCX files.")
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            raise ValueError("Unsupported file format. Use PDF or DOCX.")
    except Exception as e:
        raise ValueError(f"Failed to extract text from {filename}: {str(e)}")

def extract_skills(text: str) -> List[str]:
    """Extract skills from text with or without spacy"""
    if SPACY_AVAILABLE and nlp:
        # Use spacy for better extraction
        doc = nlp(text.lower())
        skills_keywords = [
            "python", "javascript", "react", "node.js", "sql", "aws", "docker",
            "java", "c++", "golang", "kubernetes", "mongodb", "postgresql",
            "machine learning", "ai", "tensorflow", "pytorch", "django", "flask",
            "fastapi", "vue", "angular", "typescript", "redis", "elasticsearch"
        ]
        
        found_skills = []
        for skill in skills_keywords:
            if skill in text.lower():
                found_skills.append(skill)
        
        return found_skills
    else:
        # Fallback to simple keyword matching
        skills_keywords = [
            "python", "javascript", "react", "node.js", "sql", "aws", "docker",
            "java", "c++", "golang", "machine learning", "ai"
        ]
        
        found_skills = []
        text_lower = text.lower()
        for skill in skills_keywords:
            if skill in text_lower:
                found_skills.append(skill)
        
        return found_skills

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information from text"""
    contact_info = {}
    
    # Extract email
    email_match = re.search(EMAIL_REGEX, text)
    contact_info['email'] = email_match.group(0) if email_match else ""
    
    # Extract phone
    phone_match = re.search(PHONE_REGEX, text)
    contact_info['phone'] = phone_match.group(0) if phone_match else ""
    
    # Extract LinkedIn
    linkedin_match = re.search(LINKEDIN_REGEX, text)
    contact_info['linkedin'] = linkedin_match.group(0) if linkedin_match else ""
    
    # Extract GitHub
    github_match = re.search(GITHUB_REGEX, text)
    contact_info['github'] = github_match.group(0) if github_match else ""
    
    return contact_info

def extract_name(text: str) -> str:
    """Extract name from resume text"""
    lines = text.split('\n')
    for line in lines[:5]:  # Check first 5 lines
        line = line.strip()
        if line and len(line.split()) <= 4 and len(line) > 2:
            # Simple heuristic: name is usually in first few lines, 1-4 words
            if not re.search(EMAIL_REGEX, line) and not re.search(PHONE_REGEX, line):
                return line
    return "Unknown"

def parse_resume(file: BytesIO, filename: str) -> Dict:
    """Main function to parse resume and extract all information"""
    try:
        # Extract text
        text = extract_text(file, filename)
        
        # Extract various components
        contact_info = extract_contact_info(text)
        skills = extract_skills(text)
        name = extract_name(text)
        
        return {
            "name": name,
            "email": contact_info.get('email', ''),
            "phone": contact_info.get('phone', ''),
            "linkedin": contact_info.get('linkedin', ''),
            "github": contact_info.get('github', ''),
            "skills": skills,
            "raw_text_preview": text[:500],  # First 500 characters
            "work_experience": [],  # Placeholder for now
            "education": [],  # Placeholder for now
            "spacy_available": SPACY_AVAILABLE,
            "docx_available": DOCX_AVAILABLE
        }
        
    except Exception as e:
        raise Exception(f"Resume parsing failed: {str(e)}")
